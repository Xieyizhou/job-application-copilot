"""Local extraction of Personal candidate documents into canonical Markdown."""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from docx import Document


SUPPORTED_CANDIDATE_EXTENSIONS = {".md", ".txt", ".docx", ".pdf"}
SCANNED_PDF_MESSAGE = (
    "No usable text was found in this PDF. It may be a scanned document. "
    "Please upload a text-based PDF, DOCX, Markdown, or TXT file."
)


class CandidateDocumentError(ValueError):
    """Raised when a candidate upload cannot be safely read locally."""


@dataclass(frozen=True)
class CandidateDocumentResult:
    """Canonical Markdown and non-sensitive extraction details."""

    markdown: str
    original_extension: str
    extraction_method: str
    page_count: int | None = None
    warnings: tuple[str, ...] = ()


def parse_candidate_document(filename: str, content: bytes) -> CandidateDocumentResult:
    """Parse a supported local upload without trusting its supplied path."""
    extension = _validated_extension(filename)
    if not content:
        raise CandidateDocumentError("Candidate source files must not be empty.")

    if extension in {".md", ".txt"}:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise CandidateDocumentError("Candidate Markdown and text files must use UTF-8 encoding.") from error
        return CandidateDocumentResult(
            markdown=_normalize_markdown(text),
            original_extension=extension,
            extraction_method="markdown" if extension == ".md" else "plain_text",
        )
    if extension == ".docx":
        return CandidateDocumentResult(
            markdown=_extract_docx(content),
            original_extension=extension,
            extraction_method="docx",
        )
    return _extract_pdf(content)


def _validated_extension(filename: str) -> str:
    if not isinstance(filename, str) or not filename or filename.strip() != filename:
        raise CandidateDocumentError("Candidate upload filename is invalid.")
    if any(ord(character) < 32 for character in filename):
        raise CandidateDocumentError("Candidate upload filename is invalid.")
    if any(separator in filename for separator in ("/", "\\", "\u2215", "\u2044", "\uff0f", "\uff3c")):
        raise CandidateDocumentError("Candidate upload filename must not contain a path.")
    if ":" in filename or filename in {".", ".."}:
        raise CandidateDocumentError("Candidate upload filename must not contain a path.")
    extension_match = re.search(r"(\.[A-Za-z0-9]+)$", filename)
    extension = extension_match.group(1).lower() if extension_match else ""
    if extension not in SUPPORTED_CANDIDATE_EXTENSIONS:
        raise CandidateDocumentError("Unsupported candidate source. Use PDF, DOCX, Markdown, or TXT.")
    return extension


def _normalize_markdown(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").lstrip("\ufeff")
    normalized = "\n".join(line.rstrip() for line in normalized.split("\n"))
    normalized = re.sub(r"\n[ \t]*\n(?:[ \t]*\n)+", "\n\n", normalized).strip()
    if not normalized or not re.search(r"\S", normalized):
        raise CandidateDocumentError("No usable candidate text was found in this file.")
    return normalized + "\n"


def _extract_docx(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except Exception as error:  # python-docx exposes several ZIP/XML failure types.
        raise CandidateDocumentError("This DOCX file could not be read. Please upload a valid DOCX file.") from error

    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        heading_match = re.search(r"heading\s*([1-3])", style_name)
        if heading_match:
            lines.append("#" * int(heading_match.group(1)) + " " + text)
        elif "list bullet" in style_name or "list number" in style_name:
            lines.append("- " + text)
        else:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            useful_cells = [cell for cell in cells if cell]
            if useful_cells:
                lines.append(" | ".join(useful_cells))

    return _normalize_markdown("\n\n".join(lines))


def _extract_pdf(content: bytes) -> CandidateDocumentResult:
    try:
        import pymupdf
    except ImportError as error:
        raise CandidateDocumentError("PDF extraction is unavailable in this installation.") from error

    try:
        document = pymupdf.open(stream=content, filetype="pdf")
    except Exception as error:
        raise CandidateDocumentError("This PDF file could not be read. Please upload a valid text-based PDF.") from error

    try:
        if document.needs_pass:
            raise CandidateDocumentError("Password-protected PDFs are not supported.")
        page_count = document.page_count
        page_text = [page.get_text("text") for page in document]
    except CandidateDocumentError:
        raise
    except Exception as error:
        raise CandidateDocumentError("This PDF file could not be read. Please upload a valid text-based PDF.") from error
    finally:
        document.close()

    combined = "\n\n".join(page_text)
    try:
        markdown = _normalize_markdown(combined)
    except CandidateDocumentError as error:
        raise CandidateDocumentError(SCANNED_PDF_MESSAGE) from error
    return CandidateDocumentResult(
        markdown=markdown,
        original_extension=".pdf",
        extraction_method="pdf_text",
        page_count=page_count,
    )
