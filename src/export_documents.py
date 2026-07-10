"""Export employer-facing Markdown application materials to Word DOCX."""

from __future__ import annotations

import argparse
import re
import warnings
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from workspace import Workspace, WorkspaceError, generic_cover_letter_template, personal_workspace


PROJECT_ROOT = Path(__file__).resolve().parents[1]
USER_NAME = "Candidate Name"
USER_EMAIL = "candidate@example.com"
USER_LOCATION = "Example City"
USER_LINKEDIN = "https://www.linkedin.com/in/example-profile/"


FORBIDDEN_EMPLOYER_PHRASES = [
    "YOUR NAME",
    "[Date]",
    "[Name of employer]",
    "[Hiring manager's name]",
    "[Your name]",
    "[Your signature]",
    "AI-generated",
    "Tailoring Notes",
    "internal notes",
    "undergraduate graduate",
    "master's",
    "Master's",
    "MS",
    "M.S.",
    "# Cover Letter",
]


def validate_employer_content(label: str, markdown_text: str) -> list[str]:
    """Return forbidden phrases found in employer-facing Markdown."""
    warnings = []

    for phrase in FORBIDDEN_EMPLOYER_PHRASES:
        if phrase in markdown_text:
            warnings.append(f"{label}: found forbidden phrase '{phrase}'")

    for line in markdown_text.splitlines():
        if line.strip() == "Cover Letter":
            warnings.append(f"{label}: found forbidden standalone title 'Cover Letter'")

    return warnings


def clean_duplicated_punctuation(text: str) -> str:
    """Fix common duplicated periods before writing employer-facing content."""
    replacements = {
        "Pte..": "Pte.",
        "Ltd..": "Ltd.",
        "Inc..": "Inc.",
        "Company..": "Company.",
    }

    for bad_text, clean_text in replacements.items():
        text = text.replace(bad_text, clean_text)

    return text


def configure_document_styles(document: Document) -> None:
    """Apply cover-letter-safe page settings while preserving template styles."""
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.0


def configure_resume_document_styles(document: Document) -> None:
    """Apply compact resume-safe page and font settings."""
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(10.5)
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.line_spacing = 1.0

    for style_name, font_size in [("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(font_size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)

    if "List Bullet" in [style.name for style in document.styles if style.type == WD_STYLE_TYPE.PARAGRAPH]:
        bullet_style = document.styles["List Bullet"]
        bullet_style.font.name = "Calibri"
        bullet_style.font.size = Pt(10.5)
        bullet_style.paragraph_format.space_after = Pt(2)


def clear_document_body(document: Document) -> None:
    """Remove template placeholder body content while keeping styles/margins."""
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def clean_cover_letter_parts(markdown_text: str) -> list[str]:
    """Return body paragraphs without Markdown headings, greeting, or signature."""
    lines = clean_duplicated_punctuation(markdown_text).splitlines()
    paragraph_lines = []
    paragraphs = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            flush_paragraph_text(paragraphs, paragraph_lines)
            paragraph_lines = []
            continue

        if line.startswith("#"):
            continue
        if line == "Cover Letter":
            continue
        if line == "Dear Hiring Team,":
            continue
        if line == "Sincerely,":
            continue
        if line == USER_NAME:
            continue

        paragraph_lines.append(line)

    flush_paragraph_text(paragraphs, paragraph_lines)
    return [clean_duplicated_punctuation(paragraph) for paragraph in paragraphs]


def flush_paragraph_text(paragraphs: list[str], paragraph_lines: list[str]) -> None:
    """Join accumulated Markdown lines into one paragraph string."""
    if not paragraph_lines:
        return
    paragraphs.append(" ".join(paragraph_lines))


def add_paragraph(
    document: Document,
    text: str = "",
    *,
    bold: bool = False,
    center: bool = False,
    font_size: float = 12,
    space_after: float = 6,
) -> None:
    """Add one explicitly formatted paragraph."""
    text = clean_duplicated_punctuation(text)
    paragraph = document.add_paragraph()
    if center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)


def clean_markdown_inline(text: str) -> str:
    """Remove simple Markdown emphasis/link markers for DOCX text."""
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return clean_duplicated_punctuation(text.strip())


def add_resume_paragraph(document: Document, text: str, *, style: str | None = None) -> None:
    """Add a compact resume paragraph with optional built-in style."""
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(2 if style == "List Bullet" else 3)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(clean_markdown_inline(text))
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)


def add_resume_markdown_to_document(document: Document, markdown_text: str) -> None:
    """Convert simple Markdown resume content into a Word document."""
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            add_resume_paragraph(document, " ".join(paragraph_lines))
            paragraph_lines = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            flush_paragraph()
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            heading_text = clean_markdown_inline(heading_match.group(2))
            document.add_heading(heading_text, level=level)
            continue

        bullet_match = re.match(r"^[-*]\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            add_resume_paragraph(document, bullet_match.group(1), style="List Bullet")
            continue

        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", line):
            flush_paragraph()
            continue

        paragraph_lines.append(line)

    flush_paragraph()


def add_cover_letter_to_template(
    document: Document,
    markdown_text: str,
    metadata: dict[str, str],
) -> None:
    """Write contact, employer metadata, and cover letter body into template."""
    body_paragraphs = clean_cover_letter_parts(markdown_text)

    add_paragraph(document, USER_NAME, bold=True, center=True, font_size=16, space_after=2)
    add_paragraph(
        document,
        f"{USER_EMAIL} | {USER_LOCATION} | {USER_LINKEDIN}",
        center=True,
        font_size=10.5,
        space_after=12,
    )

    today = datetime.now()
    add_paragraph(document, f"{today.strftime('%B')} {today.day}, {today.year}", space_after=6)

    if metadata.get("company"):
        add_paragraph(document, metadata["company"])
    if metadata.get("location"):
        add_paragraph(document, metadata["location"], space_after=12)
    else:
        add_paragraph(document, "", space_after=6)

    add_paragraph(document, "Dear Hiring Team,", space_after=6)

    for paragraph_text in body_paragraphs:
        add_paragraph(document, paragraph_text, space_after=6)

    add_paragraph(document, "Sincerely,", space_after=6)
    add_paragraph(document, USER_NAME, space_after=0)

def final_document_text(document: Document) -> str:
    """Return all non-empty paragraph text from the DOCX object."""
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())


def parse_job_metadata_from_package(package_dir: Path) -> dict[str, str]:
    """Find company, role, and location by following analysis.md to the job file."""
    metadata = {"company": "", "role": "", "location": ""}
    analysis_path = package_dir / "analysis.md"
    if not analysis_path.exists():
        return metadata

    analysis_text = analysis_path.read_text(encoding="utf-8")
    match = re.search(r"Job description file: `(.+?)`", analysis_text)
    if not match:
        return metadata

    job_path = Path(match.group(1))
    if not job_path.exists():
        return metadata

    job_text = job_path.read_text(encoding="utf-8")
    for line in job_text.splitlines():
        for key, label in [("company", "Company"), ("role", "Role"), ("location", "Location")]:
            prefix = f"{label}:"
            if line.lower().startswith(prefix.lower()):
                value = line.split(":", 1)[1].strip()
                if value and value.lower() != "not provided":
                    metadata[key] = value

    return metadata


def export_cover_letter_to_docx(
    markdown_path: Path,
    docx_path: Path,
    metadata: dict[str, str],
    template_path: Path,
) -> list[str]:
    """Validate and export cover_letter.md to cover_letter.docx."""
    markdown_text = clean_duplicated_punctuation(markdown_path.read_text(encoding="utf-8"))
    warnings = validate_employer_content(markdown_path.name, markdown_text)

    document = Document(template_path)
    configure_document_styles(document)
    clear_document_body(document)
    add_cover_letter_to_template(document, markdown_text, metadata)
    warnings.extend(validate_employer_content("cover_letter.docx", final_document_text(document)))
    document.save(docx_path)

    return warnings


def export_resume_to_docx(markdown_path: Path, docx_path: Path) -> list[str]:
    """Export tailored_resume.md to a simple resume DOCX."""
    markdown_text = clean_duplicated_punctuation(markdown_path.read_text(encoding="utf-8"))
    warnings = validate_employer_content(markdown_path.name, markdown_text)

    document = Document()
    configure_resume_document_styles(document)
    add_resume_markdown_to_document(document, markdown_text)
    warnings.extend(validate_employer_content(docx_path.name, final_document_text(document)))
    document.save(docx_path)
    return warnings


def export_application_package(
    package_dir: Path,
    workspace: Workspace,
) -> tuple[Path, Path, list[str]]:
    """Export employer-facing resume and cover letter files from one package folder."""
    workspace.require_writable()
    resume_markdown = package_dir / "tailored_resume.md"
    cover_letter_markdown = package_dir / "cover_letter.md"
    template_path = generic_cover_letter_template(workspace)

    if not template_path.exists():
        raise FileNotFoundError(f"Missing cover letter template: {template_path}")
    if not resume_markdown.exists():
        raise FileNotFoundError(f"Missing required file: {resume_markdown}")
    if not cover_letter_markdown.exists():
        raise FileNotFoundError(f"Missing required file: {cover_letter_markdown}")

    resume_docx = package_dir / "tailored_resume.docx"
    cover_letter_docx = package_dir / "cover_letter.docx"
    metadata = parse_job_metadata_from_package(package_dir)
    export_warnings = export_resume_to_docx(resume_markdown, resume_docx)
    export_warnings.extend(
        export_cover_letter_to_docx(cover_letter_markdown, cover_letter_docx, metadata, template_path)
    )
    return resume_docx, cover_letter_docx, export_warnings


def parse_args() -> argparse.Namespace:
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(
        description="Export tailored_resume.md and cover_letter.md to DOCX."
    )
    parser.add_argument("package_dir", help="Application package folder containing Markdown files.")
    return parser.parse_args()


def main() -> None:
    """Command-line entry point."""
    args = parse_args()
    package_dir = Path(args.package_dir).expanduser()

    if not package_dir.is_absolute():
        package_dir = Path.cwd() / package_dir

    if not package_dir.exists() or not package_dir.is_dir():
        raise FileNotFoundError(f"Package folder was not found: {package_dir}")

    workspace = personal_workspace()
    try:
        workspace.require_ready()
    except WorkspaceError as error:
        raise SystemExit(str(error)) from None
    template_found = generic_cover_letter_template(workspace).exists()
    resume_md_found = (package_dir / "tailored_resume.md").exists()
    cover_letter_md_found = (package_dir / "cover_letter.md").exists()
    resume_docx, cover_letter_docx, warnings = export_application_package(package_dir, workspace)

    print(f"Package folder: {package_dir}")
    print(f"Template file found: {template_found}")
    print(f"tailored_resume.md found: {resume_md_found}")
    print(f"cover_letter.md found: {cover_letter_md_found}")
    print(f"Created resume DOCX: {resume_docx}")
    print(f"Created cover letter DOCX: {cover_letter_docx}")

    if warnings:
        print("Validation warnings:")
        for warning in warnings:
            print(f"- {warning}")
    else:
        print("Validation warnings: none")

    print("Reminder: manually review DOCX files before submitting.")


if __name__ == "__main__":
    main()
