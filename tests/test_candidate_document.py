"""Sanitized candidate upload parsing and canonical-workspace coverage."""

from __future__ import annotations

import io
import json
import sys
import tempfile
import unittest
from pathlib import Path

import pymupdf
from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

from analyze_job import analyze_job
from candidate_document import CandidateDocumentError, SCANNED_PDF_MESSAGE, parse_candidate_document
from workspace import WorkspaceError, initialize_personal_workspace, personal_workspace


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Fictional Candidate", level=1)
    document.add_paragraph("Python and SQL analyst")
    document.add_paragraph("Sanitized project result", style="List Bullet")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_pdf(page_text: list[str]) -> bytes:
    document = pymupdf.open()
    for text in page_text:
        page = document.new_page()
        if text:
            page.insert_text((72, 72), text)
    result = document.tobytes()
    document.close()
    return result


class CandidateDocumentTests(unittest.TestCase):
    def test_markdown_and_utf8_text_are_normalized(self) -> None:
        markdown = parse_candidate_document("candidate.MD", b"# Candidate\r\n\r\n\r\n- Python\r\n")
        text = parse_candidate_document("candidate.txt", "\ufeffFictional Candidate\r\nSQL\r\n".encode())
        self.assertEqual(markdown.markdown, "# Candidate\n\n- Python\n")
        self.assertEqual(markdown.extraction_method, "markdown")
        self.assertEqual(text.markdown, "Fictional Candidate\nSQL\n")
        self.assertEqual(text.extraction_method, "plain_text")

    def test_invalid_text_and_unsafe_names_are_rejected(self) -> None:
        for filename, content in [
            ("empty.txt", b""),
            ("bad.txt", b"\xff"),
            ("../candidate.md", b"# Candidate"),
            (r"..\\candidate.md", b"# Candidate"),
            ("C:resume.md", b"# Candidate"),
        ]:
            with self.assertRaises(CandidateDocumentError):
                parse_candidate_document(filename, content)

    def test_docx_extracts_headings_paragraphs_lists_and_tables(self) -> None:
        result = parse_candidate_document("resume.docx", make_docx())
        self.assertEqual(result.extraction_method, "docx")
        self.assertIn("# Fictional Candidate", result.markdown)
        self.assertIn("Python and SQL analyst", result.markdown)
        self.assertIn("- Sanitized project result", result.markdown)
        self.assertIn("Skill | Python", result.markdown)
        empty_document = Document()
        empty_buffer = io.BytesIO()
        empty_document.save(empty_buffer)
        with self.assertRaises(CandidateDocumentError):
            parse_candidate_document("empty.docx", empty_buffer.getvalue())
        with self.assertRaises(CandidateDocumentError):
            parse_candidate_document("corrupt.docx", b"not a document")

    def test_pdf_extracts_pages_in_order_and_rejects_unusable_inputs(self) -> None:
        result = parse_candidate_document("resume.PDF", make_pdf(["First page Python", "Second page SQL"]))
        self.assertEqual(result.extraction_method, "pdf_text")
        self.assertEqual(result.page_count, 2)
        self.assertLess(result.markdown.index("First page Python"), result.markdown.index("Second page SQL"))
        with self.assertRaisesRegex(CandidateDocumentError, SCANNED_PDF_MESSAGE):
            parse_candidate_document("scanned.pdf", make_pdf([""]))
        with self.assertRaises(CandidateDocumentError):
            parse_candidate_document("corrupt.pdf", b"not a pdf")

        encrypted = pymupdf.open()
        encrypted.new_page().insert_text((72, 72), "Private text")
        encrypted_bytes = encrypted.tobytes(
            encryption=pymupdf.PDF_ENCRYPT_AES_256,
            owner_pw="owner",
            user_pw="user",
        )
        encrypted.close()
        with self.assertRaisesRegex(CandidateDocumentError, "Password-protected"):
            parse_candidate_document("protected.pdf", encrypted_bytes)

    def test_safe_replacement_uses_generic_files_and_preserves_workspace_assets(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_dir:
            root = Path(temporary_dir) / "local_workspace"
            workspace = initialize_personal_workspace(
                "first.md",
                b"# First\nPython",
                experience_bank=("experience.yaml", b"items: []\n"),
                cover_letter_template=("template.docx", b"template placeholder"),
                root=root,
            )
            (workspace.jobs_dir / "saved.md").write_text("sanitized job", encoding="utf-8")
            workspace.generated_dir.mkdir(exist_ok=True)
            (workspace.generated_dir / "output.md").write_text("sanitized output", encoding="utf-8")
            assert workspace.tracker_database_path is not None
            workspace.tracker_database_path.write_bytes(b"sqlite placeholder")
            assert workspace.cover_letter_template_path is not None
            assert workspace.experience_bank_path is not None
            template = workspace.cover_letter_template_path
            experience = workspace.experience_bank_path
            manifest_path = root / "workspace.json"
            original_manifest = manifest_path.read_text(encoding="utf-8")
            original_source = workspace.resume_source_path.read_text(encoding="utf-8") if workspace.resume_source_path else ""

            with self.assertRaises(WorkspaceError):
                initialize_personal_workspace("bad.pdf", b"not a pdf", root=root)
            self.assertEqual(manifest_path.read_text(encoding="utf-8"), original_manifest)
            self.assertEqual(workspace.resume_source_path.read_text(encoding="utf-8"), original_source)
            self.assertTrue(personal_workspace(root).ready)

            updated = initialize_personal_workspace("person.DOCX", make_docx(), root=root)
            self.assertTrue(updated.ready)
            self.assertEqual(updated.resume_source_path, root.resolve() / "candidate" / "candidate_source.md")
            self.assertTrue((root / "candidate" / "original_resume.docx").is_file())
            self.assertFalse((root / "candidate" / "original_resume.md").exists())
            self.assertTrue((workspace.jobs_dir / "saved.md").is_file())
            self.assertTrue((workspace.generated_dir / "output.md").is_file())
            self.assertTrue(workspace.tracker_database_path.is_file())
            self.assertTrue(template.is_file())
            self.assertTrue(experience.is_file())
            self.assertEqual(updated.cover_letter_template_path, template)
            self.assertEqual(updated.experience_bank_path, experience)
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(manifest["resume_source"], "candidate/candidate_source.md")
            self.assertEqual(manifest["candidate_original_extension"], ".docx")
            self.assertNotIn("person.DOCX", manifest_path.read_text(encoding="utf-8"))

    def test_docx_and_pdf_workspaces_feed_canonical_markdown_to_analysis(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_dir:
            root = Path(temporary_dir) / "local_workspace"
            for filename, content in [
                ("candidate.docx", make_docx()),
                ("candidate.pdf", make_pdf(["Fictional Candidate Python SQL"])),
            ]:
                workspace = initialize_personal_workspace(filename, content, root=root)
                job_path = workspace.jobs_dir / "job.md"
                job_path.write_text("# Role\nRequirements: Python and SQL", encoding="utf-8")
                report, _ = analyze_job(job_path, workspace)
                self.assertIn("Python", report)
                self.assertEqual(workspace.resume_source_path.name, "candidate_source.md")
                self.assertTrue(personal_workspace(root).ready)



if __name__ == "__main__":
    unittest.main()
